"""
render.py — three renderers over the same field results.

All three consume the identical record list and the same output template, so a
field always reads the same way whichever format you export. Each keeps the two
V1 safety properties:

  * every field carries a visible "Source:" citation
  * anything not found is loud (red in .docx/.pdf, bold marker in .md) and
    listed in a "Things to check" block at the very top

PDF is generated directly with ReportLab rather than by converting the .docx.
Converting would need Word or LibreOffice driving a conversion on the machine -
an install and an automation dependency this environment can't assume - whereas
ReportLab is a pure-Python writer that works offline with no external process.
"""

import datetime
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (SimpleDocTemplate, Paragraph as RLPara, Spacer,
                                Table as RLTable, TableStyle)

GREY = RGBColor(0x70, 0x70, 0x70)
RED = RGBColor(0xA3, 0x20, 0x20)
NOT_FOUND = "NOT FOUND — please verify manually"


def _ordered(records, template):
    """
    Walk the template and pair each section with its field result.
    Fields present in the rule set but absent from the template are appended at
    the end, so adding a rule never makes its output silently disappear.
    """
    by_name = {r["name"]: r for r in records}
    used, items = set(), []
    for section in (template or {}).get("sections", []):
        if section.get("type") == "heading":
            items.append(("heading", section.get("text", "")))
        else:
            fname = section.get("field")
            if fname in by_name:
                items.append(("field", by_name[fname]))
                used.add(fname)
    for r in records:
        if r["name"] not in used:
            items.append(("field", r))
    return items


def _flagged(records):
    return [r for r in records if not r["found"] or r.get("note")]


# ==========================================================================
# WORD
# ==========================================================================

def _borders(table):
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "CCCCCC")
        b.append(el)
    table._tbl.tblPr.append(b)


def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_table(doc, rows):
    if not rows:
        return
    n = max(len(r) for r in rows) or 1
    t = doc.add_table(rows=0, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    _borders(t)
    w = Inches(6.5 / n)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(n):
            cells[ci].width = w
            run = cells[ci].paragraphs[0].add_run(row[ci] if ci < len(row) else "")
            run.font.size = Pt(9.5)
            if ri == 0:
                run.bold = True
                _shade(cells[ci], "E8EEF4")


def _docx_value(doc, rec):
    kind, value = rec["kind"], rec["value"]
    if kind == "text":
        doc.add_paragraph().add_run(str(value))
    elif kind == "table":
        _docx_table(doc, value.get("rows", []))
    elif kind == "block":
        for b in value:
            if b["type"] == "paragraph":
                p = doc.add_paragraph()
                if b.get("runs"):
                    for r in b["runs"]:
                        run = p.add_run(r["text"])
                        run.bold = r.get("bold", False)
                        run.italic = r.get("italic", False)
                else:
                    p.add_run(b["text"])
            else:
                _docx_table(doc, b["rows"])


def render_docx(records, template, out_path, meta):
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

    doc.add_heading((template or {}).get("title", "Extraction Summary"), level=0)
    m = doc.add_paragraph()
    r = m.add_run(f"Generated {meta['generated']}\nSource files: {meta['files']}")
    r.font.size = Pt(9); r.font.color.rgb = GREY

    flags = _flagged(records)
    if flags:
        h = doc.add_heading("Things to check", level=1)
        h.runs[0].font.color.rgb = RED
        t = doc.add_table(rows=1, cols=1); _borders(t)
        cell = t.rows[0].cells[0]; _shade(cell, "FBE9E7")
        cell.paragraphs[0].add_run(
            "Verify the following before relying on this summary:").bold = True
        for f in flags:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            p.add_run(f"•  {f['name']}: "
                      f"{f.get('note') or 'Not found in the documents.'}"
                      ).font.size = Pt(10)
    doc.add_paragraph()

    for kind, item in _ordered(records, template):
        if kind == "heading":
            doc.add_heading(item, level=1)
            continue
        rec = item
        doc.add_heading(rec["name"], level=2)
        if not rec["found"]:
            run = doc.add_paragraph().add_run(rec.get("note") or NOT_FOUND)
            run.font.color.rgb = RED; run.bold = True
        else:
            _docx_value(doc, rec)
            if rec.get("note"):
                n = doc.add_paragraph().add_run(f"Note: {rec['note']}")
                n.font.size = Pt(9); n.font.color.rgb = GREY; n.italic = True
        c = doc.add_paragraph().add_run(f"Source: {rec['citation']}")
        c.font.size = Pt(9); c.font.color.rgb = GREY; c.italic = True
    doc.save(out_path)
    return out_path


# ==========================================================================
# MARKDOWN
# ==========================================================================

def _md_table(rows):
    if not rows:
        return ""
    n = max(len(r) for r in rows) or 1
    norm = [list(r) + [""] * (n - len(r)) for r in rows]
    esc = lambda s: str(s).replace("|", "\\|").replace("\n", " ")
    out = ["| " + " | ".join(esc(c) for c in norm[0]) + " |",
           "|" + "|".join([" --- "] * n) + "|"]
    for row in norm[1:]:
        out.append("| " + " | ".join(esc(c) for c in row) + " |")
    return "\n".join(out)


def render_md(records, template, out_path, meta):
    L = [f"# {(template or {}).get('title', 'Extraction Summary')}", "",
         f"*Generated {meta['generated']}*  ", f"*Source files: {meta['files']}*", ""]

    flags = _flagged(records)
    if flags:
        L += ["> **Things to check**", ">",
              "> Verify the following before relying on this summary:", ">"]
        for f in flags:
            L.append(f"> - **{f['name']}**: "
                     f"{f.get('note') or 'Not found in the documents.'}")
        L.append("")

    for kind, item in _ordered(records, template):
        if kind == "heading":
            L += [f"## {item}", ""]
            continue
        rec = item
        L += [f"### {rec['name']}", ""]
        if not rec["found"]:
            L += [f"**{rec.get('note') or NOT_FOUND}**", ""]
        else:
            if rec["kind"] == "text":
                L += [str(rec["value"]), ""]
            elif rec["kind"] == "table":
                L += [_md_table(rec["value"].get("rows", [])), ""]
            else:
                for b in rec["value"]:
                    if b["type"] == "paragraph":
                        L += [b["text"], ""]
                    else:
                        L += [_md_table(b["rows"]), ""]
            if rec.get("note"):
                L += [f"*Note: {rec['note']}*", ""]
        L += [f"`Source: {rec['citation']}`", ""]

    with open(out_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(L))
    return out_path


# ==========================================================================
# PDF
# ==========================================================================

def _esc(s):
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _est_row_height(row, col_width, font_size=9, leading=12):
    """
    Roughly how tall a table row will be once its text wraps.

    ReportLab cannot split a table row across pages: if one row is taller than
    the printable frame, the whole build fails with a "too large on page" error.
    So we estimate first and degrade gracefully instead of crashing.
    """
    chars_per_line = max(int(col_width / (font_size * 0.5)), 8)
    tallest = 0
    for cell in row:
        text = str(cell or "")
        lines = 0
        for para in text.split("\n"):
            lines += max(1, -(-len(para) // chars_per_line))  # ceil division
        tallest = max(tallest, lines * leading)
    return tallest + 8  # cell padding


def render_pdf(records, template, out_path, meta):
    ss = getSampleStyleSheet()
    title_s = ParagraphStyle("t", parent=ss["Title"], fontSize=18, alignment=0,
                             spaceAfter=6)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=13, spaceBefore=12)
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontSize=11, spaceBefore=8)
    body = ParagraphStyle("b", parent=ss["BodyText"], fontSize=10, leading=14)
    small = ParagraphStyle("s", parent=ss["BodyText"], fontSize=8,
                           textColor=colors.HexColor("#707070"))
    warn = ParagraphStyle("w", parent=ss["BodyText"], fontSize=10,
                          textColor=colors.HexColor("#A32020"))

    story = [RLPara(_esc((template or {}).get("title", "Extraction Summary")), title_s),
             RLPara(f"Generated {_esc(meta['generated'])}<br/>"
                    f"Source files: {_esc(meta['files'])}", small),
             Spacer(1, 10)]

    flags = _flagged(records)
    if flags:
        story.append(RLPara("Things to check", h1))
        for f in flags:
            story.append(RLPara(
                f"• <b>{_esc(f['name'])}</b>: "
                f"{_esc(f.get('note') or 'Not found in the documents.')}", warn))
        story.append(Spacer(1, 8))

    # The usable height inside the page frame. A row taller than this cannot be
    # rendered as a table at all, because ReportLab never splits mid-row.
    MAX_ROW = 620

    def pdf_table(rows):
        if not rows:
            return
        n = max(len(r) for r in rows) or 1
        col_w = (6.5 * inch) / n

        # If ANY row is too tall to fit a page, render the whole table as
        # labelled paragraphs instead. Content is preserved; only the grid is
        # lost, which is far better than failing the export.
        if any(_est_row_height(r, col_w) > MAX_ROW for r in rows):
            header = rows[0] if len(rows) > 1 else []
            for row in (rows[1:] if len(rows) > 1 else rows):
                for i, cell in enumerate(row):
                    if not str(cell).strip():
                        continue
                    label = str(header[i]).strip() if i < len(header) else ""
                    text = f"<b>{_esc(label)}:</b> {_esc(cell)}" if label else _esc(cell)
                    story.append(RLPara(text, body))
                story.append(Spacer(1, 4))
            story.append(RLPara(
                "(Table shown as text: one row was too tall to fit a page.)", small))
            story.append(Spacer(1, 6))
            return

        cell_style = ParagraphStyle("cell", parent=body, fontSize=9, leading=12)
        data = [[RLPara(_esc(r[i] if i < len(r) else ""), cell_style)
                 for i in range(n)] for r in rows]
        t = RLTable(data, colWidths=[col_w] * n, repeatRows=1, splitByRow=1)
        t.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF4")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4)]))
        story.append(t)
        story.append(Spacer(1, 6))

    for kind, item in _ordered(records, template):
        if kind == "heading":
            story.append(RLPara(_esc(item), h1))
            continue
        rec = item
        story.append(RLPara(_esc(rec["name"]), h2))
        if not rec["found"]:
            story.append(RLPara(_esc(rec.get("note") or NOT_FOUND), warn))
        else:
            if rec["kind"] == "text":
                story.append(RLPara(_esc(rec["value"]), body))
            elif rec["kind"] == "table":
                pdf_table(rec["value"].get("rows", []))
            else:
                for b in rec["value"]:
                    if b["type"] == "paragraph":
                        if b["text"].strip():
                            story.append(RLPara(_esc(b["text"]), body))
                    else:
                        pdf_table(b["rows"])
            if rec.get("note"):
                story.append(RLPara(f"Note: {_esc(rec['note'])}", small))
        story.append(RLPara(f"Source: {_esc(rec['citation'])}", small))
        story.append(Spacer(1, 6))

    def _build(st):
        SimpleDocTemplate(out_path, pagesize=LETTER,
                          leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                          topMargin=0.9 * inch, bottomMargin=0.9 * inch).build(st)

    try:
        _build(story)
    except Exception:
        # Last resort: some flowable still didn't fit. Strip every table down to
        # plain paragraphs so the export always produces a usable document
        # rather than failing outright.
        flat = []
        for item in story:
            if isinstance(item, RLTable):
                try:
                    for row in item._cellvalues:
                        line = " · ".join(
                            getattr(c, "text", str(c)) for c in row
                            if getattr(c, "text", str(c)).strip())
                        if line:
                            flat.append(RLPara(line, body))
                except Exception:
                    continue
            else:
                flat.append(item)
        _build(flat)
    return out_path


# ==========================================================================
# JSON
# ==========================================================================

def render_json(records, template, out_path, meta):
    """
    Machine-readable output: every field with its value, citation and whether it
    was found. Block/table values are flattened into plain structures so the
    file is usable by any downstream tool without knowing this codebase.
    """
    import json

    def value_of(rec):
        if not rec["found"]:
            return None
        if rec["kind"] == "text":
            return str(rec["value"])
        if rec["kind"] == "table":
            return {"type": "table", "rows": rec["value"].get("rows", [])}
        out = []
        for b in rec["value"]:
            if b["type"] == "paragraph":
                out.append({"type": "paragraph", "text": b["text"]})
            else:
                out.append({"type": "table", "rows": b["rows"]})
        return out

    payload = {
        "title": (template or {}).get("title", "Extraction Summary"),
        "generated": meta["generated"],
        "source_files": meta["files"],
        "fields": [{
            "name": r["name"],
            "found": r["found"],
            "value": value_of(r),
            "citation": r["citation"],
            "rule": r.get("rule_type"),
            "note": r.get("note"),
        } for r in records],
        "needs_checking": [r["name"] for r in records if not r["found"]],
    }
    with open(out_path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, indent=2, ensure_ascii=False)
    return out_path


RENDERERS = {"docx": render_docx, "md": render_md, "pdf": render_pdf,
             "json": render_json}
